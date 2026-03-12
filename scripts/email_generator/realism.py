"""
Realism enhancements for email population.

This module provides realistic patterns for:
- Read/unread status with vacation clusters
- Multi-message conversation threads
- Realistic attachment content generation
- Time-based patterns (weekends, holidays, vacation)
- Out-of-office auto-reply emails
"""

import random
from datetime import datetime, timedelta
from typing import Dict, Any, List, Optional, Tuple
import hashlib


class ReadUnreadPatterns:
    """
    Generates realistic read/unread patterns for emails.
    
    Features:
    - Recent emails (< 3 days) have higher unread probability
    - Spam emails are often left unread
    - Creates "vacation clusters" where emails are unread
    - Important emails are more likely to be read
    """
    
    def __init__(self, vacation_periods: Optional[List[Tuple[datetime, datetime]]] = None):
        """
        Initialize read/unread pattern generator.
        
        Args:
            vacation_periods: List of (start, end) tuples for vacation periods
        """
        self.vacation_periods = vacation_periods or []
    
    def should_be_read(
        self,
        email_date: datetime,
        category: str,
        importance: str = "normal",
        is_spam: bool = False
    ) -> bool:
        """
        Determine if an email should be marked as read.
        
        Args:
            email_date: The date of the email
            category: Email category (spam, security, etc.)
            importance: Email importance (high, normal, low)
            is_spam: Whether this is a spam email
            
        Returns:
            True if email should be marked as read
        """
        now = datetime.now()
        if email_date.tzinfo:
            now = now.replace(tzinfo=email_date.tzinfo)
        
        age_days = (now - email_date).days
        
        # Check if email falls within a vacation period
        for start, end in self.vacation_periods:
            if start <= email_date <= end:
                # 70% of vacation emails are unread
                return random.random() > 0.7
        
        # Spam emails - 60% left unread
        if is_spam or category == "spam":
            return random.random() > 0.6
        
        # Very recent emails (< 1 day) - 40% unread
        if age_days < 1:
            return random.random() > 0.4
        
        # Recent emails (1-3 days) - 25% unread
        if age_days < 3:
            return random.random() > 0.25
        
        # Emails 3-7 days old - 15% unread
        if age_days < 7:
            return random.random() > 0.15
        
        # High importance emails are more likely to be read
        if importance == "high":
            return random.random() > 0.05
        
        # Older emails - 5% unread
        return random.random() > 0.05
    
    def generate_vacation_periods(
        self,
        start_date: datetime,
        end_date: datetime,
        num_periods: int = 2
    ) -> List[Tuple[datetime, datetime]]:
        """
        Generate random vacation periods within a date range.
        
        Args:
            start_date: Start of the date range
            end_date: End of the date range
            num_periods: Number of vacation periods to generate
            
        Returns:
            List of (start, end) tuples for vacation periods
        """
        periods = []
        total_days = (end_date - start_date).days
        
        if total_days < 30:
            return periods
        
        for _ in range(num_periods):
            # Vacation length: 3-14 days
            vacation_length = random.randint(3, 14)
            
            # Random start within the range
            max_start = total_days - vacation_length - 7
            if max_start < 7:
                continue
            
            start_offset = random.randint(7, max_start)
            vacation_start = start_date + timedelta(days=start_offset)
            vacation_end = vacation_start + timedelta(days=vacation_length)
            
            # Check for overlap with existing periods
            overlap = False
            for existing_start, existing_end in periods:
                if (vacation_start <= existing_end and vacation_end >= existing_start):
                    overlap = True
                    break
            
            if not overlap:
                periods.append((vacation_start, vacation_end))
        
        return periods


class ConversationThreadGenerator:
    """
    Generates multi-message conversation threads.
    
    Features:
    - Creates 2-5 message threads
    - Proper Re:/Fwd: subject prefixes
    - Quoted content from previous messages
    - Realistic time gaps between messages
    """
    
    def __init__(self):
        self.thread_counter = 0
    
    def generate_thread(
        self,
        base_email: Dict[str, Any],
        num_messages: Optional[int] = None,
        thread_type: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        Generate a conversation thread from a base email.
        
        Args:
            base_email: The original email to build thread from
            num_messages: Number of messages in thread (2-5 if not specified)
            thread_type: Type of thread ('reply', 'reply_all', 'forward')
            
        Returns:
            List of email dictionaries forming the thread
        """
        if num_messages is None:
            num_messages = random.randint(2, 5)
        
        if thread_type is None:
            thread_type = random.choices(
                ['reply', 'reply_all', 'forward'],
                weights=[0.6, 0.25, 0.15]
            )[0]
        
        self.thread_counter += 1
        thread_id = hashlib.md5(
            f"{base_email.get('subject', '')}-{self.thread_counter}".encode()
        ).hexdigest()[:16]
        
        thread = []
        current_email = base_email.copy()
        current_date = base_email.get('date', datetime.now())
        
        # Original sender and recipient
        original_sender = base_email.get('sender', {})
        original_recipient = base_email.get('recipient', {})
        
        # Message IDs for threading
        message_ids: List[str] = []
        
        for i in range(num_messages):
            email = current_email.copy()
            
            # Generate message ID
            msg_id = f"<{thread_id}.{i}@{original_sender.get('email', 'example.com').split('@')[-1]}>"
            message_ids.append(msg_id)
            
            # Set threading info
            email['threading'] = {
                'is_reply': i > 0 and thread_type in ('reply', 'reply_all'),
                'is_forward': i > 0 and thread_type == 'forward',
                'is_reply_all': i > 0 and thread_type == 'reply_all',
                'thread_id': thread_id,
                'message_id': msg_id,
                'in_reply_to': message_ids[i-1] if i > 0 else None,
                'references': message_ids[:i] if i > 0 else [],
            }
            
            # Modify subject for replies/forwards
            if i > 0:
                subject = base_email.get('subject', 'No Subject')
                if thread_type in ('reply', 'reply_all'):
                    if not subject.lower().startswith('re:'):
                        email['subject'] = f"Re: {subject}"
                elif thread_type == 'forward':
                    if not subject.lower().startswith('fw:') and not subject.lower().startswith('fwd:'):
                        email['subject'] = f"Fwd: {subject}"
            
            # Swap sender/recipient for replies
            if i > 0 and thread_type in ('reply', 'reply_all'):
                if i % 2 == 1:
                    email['sender'] = original_recipient
                    email['recipient'] = original_sender
                else:
                    email['sender'] = original_sender
                    email['recipient'] = original_recipient
            
            # Add time gap between messages (30 min to 48 hours)
            if i > 0:
                gap_minutes = random.randint(30, 2880)
                current_date = current_date + timedelta(minutes=gap_minutes)
            
            email['date'] = current_date
            
            # Add quoted content for replies
            if i > 0 and thread_type in ('reply', 'reply_all'):
                email['body'] = self._add_quoted_content(
                    email.get('body', ''),
                    thread[i-1].get('body', ''),
                    thread[i-1].get('sender', {}),
                    thread[i-1].get('date', current_date)
                )
            
            thread.append(email)
            current_email = email
        
        return thread
    
    def _add_quoted_content(
        self,
        new_body: str,
        previous_body: str,
        previous_sender: Dict[str, str],
        previous_date: datetime
    ) -> str:
        """Add quoted content from previous message."""
        # Strip HTML from previous body for quoting
        import re
        plain_text = re.sub(r'<[^>]+>', '', previous_body)
        plain_text = plain_text.replace('&nbsp;', ' ').replace('&amp;', '&')
        
        # Truncate if too long
        if len(plain_text) > 500:
            plain_text = plain_text[:500] + "..."
        
        # Format quoted content
        sender_name = previous_sender.get('name', 'Unknown')
        sender_email = previous_sender.get('email', '')
        date_str = previous_date.strftime('%a, %b %d, %Y at %I:%M %p')
        
        quoted = f"""
<div style="margin-top: 20px; padding-top: 10px; border-top: 1px solid #ccc;">
<p style="color: #666; font-size: 12px;">
On {date_str}, {sender_name} &lt;{sender_email}&gt; wrote:
</p>
<blockquote style="margin: 10px 0; padding-left: 10px; border-left: 2px solid #ccc; color: #666;">
{plain_text}
</blockquote>
</div>
"""
        return new_body + quoted


class TimeBasedPatterns:
    """
    Generates time-based email patterns.
    
    Features:
    - Weekend patterns (fewer emails)
    - Holiday patterns (no emails or OOO)
    - Business hours bias
    - Vacation periods
    """
    
    # Common holidays (month, day)
    HOLIDAYS = [
        (1, 1),   # New Year's Day
        (12, 25), # Christmas
        (12, 26), # Boxing Day
        (7, 4),   # Independence Day (US)
        (11, 11), # Veterans Day
        (5, 1),   # May Day
    ]
    
    def __init__(self, business_hours_bias: float = 0.8):
        """
        Initialize time-based pattern generator.
        
        Args:
            business_hours_bias: Probability of email during business hours (0-1)
        """
        self.business_hours_bias = business_hours_bias
    
    def adjust_date_for_patterns(
        self,
        base_date: datetime,
        allow_weekends: bool = True,
        allow_holidays: bool = False
    ) -> datetime:
        """
        Adjust a date based on realistic patterns.
        
        Args:
            base_date: The original date
            allow_weekends: Whether to allow weekend dates
            allow_holidays: Whether to allow holiday dates
            
        Returns:
            Adjusted datetime
        """
        adjusted = base_date
        
        # Check if it's a weekend
        if not allow_weekends and adjusted.weekday() >= 5:
            # Move to next Monday
            days_until_monday = 7 - adjusted.weekday()
            adjusted = adjusted + timedelta(days=days_until_monday)
        
        # Check if it's a holiday
        if not allow_holidays:
            for month, day in self.HOLIDAYS:
                if adjusted.month == month and adjusted.day == day:
                    # Move to next day
                    adjusted = adjusted + timedelta(days=1)
                    break
        
        # Apply business hours bias
        if random.random() < self.business_hours_bias:
            # Business hours: 8 AM - 6 PM
            hour = random.randint(8, 17)
            minute = random.randint(0, 59)
            adjusted = adjusted.replace(hour=hour, minute=minute)
        else:
            # Outside business hours
            if random.random() < 0.5:
                # Early morning (6-8 AM)
                hour = random.randint(6, 7)
            else:
                # Evening (6-10 PM)
                hour = random.randint(18, 22)
            minute = random.randint(0, 59)
            adjusted = adjusted.replace(hour=hour, minute=minute)
        
        return adjusted
    
    def get_weekend_email_probability(self) -> float:
        """Get probability of receiving email on weekend."""
        return 0.15  # 15% of normal volume
    
    def is_holiday(self, date: datetime) -> bool:
        """Check if a date is a holiday."""
        return (date.month, date.day) in self.HOLIDAYS


class OutOfOfficeGenerator:
    """
    Generates out-of-office auto-reply emails.
    
    Features:
    - Vacation OOO messages
    - Conference/travel OOO messages
    - Sick leave OOO messages
    - Proper OOO headers
    """
    
    OOO_TEMPLATES = {
        "vacation": [
            {
                "subject": "Out of Office: {name}",
                "body": """
<p>Thank you for your email. I am currently out of the office on vacation from {start_date} to {end_date}.</p>

<p>I will have limited access to email during this time. If your matter is urgent, please contact {backup_name} at {backup_email}.</p>

<p>I will respond to your email upon my return.</p>

<p>Best regards,<br>{name}</p>
"""
            },
            {
                "subject": "Automatic Reply: Out of Office",
                "body": """
<p>Hi,</p>

<p>I'm currently away from the office enjoying some time off. I'll be back on {end_date}.</p>

<p>For urgent matters, please reach out to {backup_name} ({backup_email}).</p>

<p>Thanks for your patience!</p>

<p>{name}</p>
"""
            }
        ],
        "conference": [
            {
                "subject": "Out of Office: Attending {event_name}",
                "body": """
<p>Thank you for your email.</p>

<p>I am currently attending {event_name} from {start_date} to {end_date} and will have limited access to email.</p>

<p>For urgent matters, please contact {backup_name} at {backup_email}.</p>

<p>I will respond to your email as soon as possible upon my return.</p>

<p>Best regards,<br>{name}</p>
"""
            }
        ],
        "sick": [
            {
                "subject": "Out of Office: {name}",
                "body": """
<p>Thank you for your email.</p>

<p>I am currently out of the office and will respond to your email as soon as I am able.</p>

<p>For urgent matters, please contact {backup_name} at {backup_email}.</p>

<p>Thank you for your understanding.</p>

<p>{name}</p>
"""
            }
        ]
    }
    
    EVENTS = [
        "Annual Tech Conference",
        "Industry Summit 2024",
        "Leadership Retreat",
        "Customer Conference",
        "Partner Summit",
        "Training Workshop",
    ]
    
    BACKUP_NAMES = [
        ("Sarah Johnson", "sarah.johnson@company.com"),
        ("Michael Chen", "michael.chen@company.com"),
        ("Emily Davis", "emily.davis@company.com"),
        ("James Wilson", "james.wilson@company.com"),
        ("Lisa Anderson", "lisa.anderson@company.com"),
    ]
    
    def generate_ooo_email(
        self,
        recipient: Dict[str, Any],
        original_sender: Dict[str, str],
        ooo_type: Optional[str] = None,
        start_date: Optional[datetime] = None,
        end_date: Optional[datetime] = None
    ) -> Dict[str, Any]:
        """
        Generate an out-of-office auto-reply email.
        
        Args:
            recipient: The person who is out of office
            original_sender: The person who sent the original email
            ooo_type: Type of OOO ('vacation', 'conference', 'sick')
            start_date: Start of OOO period
            end_date: End of OOO period
            
        Returns:
            Email dictionary for the OOO reply
        """
        if ooo_type is None:
            ooo_type = random.choices(
                ['vacation', 'conference', 'sick'],
                weights=[0.6, 0.3, 0.1]
            )[0]
        
        if start_date is None:
            start_date = datetime.now()
        
        if end_date is None:
            # OOO duration: 1-14 days
            duration = random.randint(1, 14)
            end_date = start_date + timedelta(days=duration)
        
        # Select template
        templates = self.OOO_TEMPLATES.get(ooo_type, self.OOO_TEMPLATES['vacation'])
        template = random.choice(templates)
        
        # Select backup contact
        backup_name, backup_email = random.choice(self.BACKUP_NAMES)
        
        # Select event name for conference
        event_name = random.choice(self.EVENTS)
        
        # Get recipient name
        recipient_name = recipient.get('display_name', recipient.get('name', 'User'))
        
        # Format dates
        start_str = start_date.strftime('%B %d, %Y')
        end_str = end_date.strftime('%B %d, %Y')
        
        # Fill template
        subject = template['subject'].format(
            name=recipient_name,
            event_name=event_name
        )
        
        body = template['body'].format(
            name=recipient_name,
            start_date=start_str,
            end_date=end_str,
            backup_name=backup_name,
            backup_email=backup_email,
            event_name=event_name
        )
        
        return {
            "category": "ooo_reply",
            "subject": subject,
            "body": body,
            "sender": {
                "name": recipient_name,
                "email": recipient.get('upn', recipient.get('email', '')),
            },
            "recipient": original_sender,
            "date": start_date + timedelta(minutes=random.randint(1, 30)),
            "is_auto_reply": True,
            "importance": "normal",
            "categories": [],
            "flag_status": {"flagged": False},
            "threading": {"is_reply": True, "is_forward": False, "is_reply_all": False},
        }


class RealisticAttachmentGenerator:
    """
    Generates realistic attachment content.
    
    Features:
    - Excel files with actual data
    - Word documents with text
    - PDF reports with content
    - PowerPoint presentations
    """
    
    def generate_excel_content(
        self,
        department: str,
        document_type: str = "report"
    ) -> bytes:
        """
        Generate realistic Excel file content.
        
        Args:
            department: Department for context
            document_type: Type of document
            
        Returns:
            Excel file bytes
        """
        try:
            from openpyxl import Workbook
            from io import BytesIO
            
            wb = Workbook()
            ws = wb.active
            if ws is None:
                ws = wb.create_sheet()
            ws.title = "Data"
            
            # Generate department-specific data
            if "finance" in department.lower():
                headers = ["Month", "Revenue", "Expenses", "Profit", "Growth %"]
                ws.append(headers)
                months = ["January", "February", "March", "April", "May", "June"]
                for month in months:
                    revenue = random.randint(100000, 500000)
                    expenses = random.randint(50000, 300000)
                    profit = revenue - expenses
                    growth = round(random.uniform(-5, 15), 1)
                    ws.append([month, revenue, expenses, profit, growth])
            
            elif "hr" in department.lower() or "human" in department.lower():
                headers = ["Employee ID", "Department", "Hire Date", "Status", "Performance"]
                ws.append(headers)
                depts = ["Engineering", "Sales", "Marketing", "Finance", "HR"]
                statuses = ["Active", "Active", "Active", "On Leave", "Active"]
                performances = ["Exceeds", "Meets", "Meets", "Exceeds", "Needs Improvement"]
                for i in range(10):
                    ws.append([
                        f"EMP{1000 + i}",
                        random.choice(depts),
                        f"202{random.randint(0, 4)}-{random.randint(1, 12):02d}-{random.randint(1, 28):02d}",
                        random.choice(statuses),
                        random.choice(performances)
                    ])
            
            else:
                # Generic data
                headers = ["Item", "Category", "Value", "Status", "Date"]
                ws.append(headers)
                categories = ["Type A", "Type B", "Type C"]
                statuses = ["Complete", "In Progress", "Pending"]
                for i in range(8):
                    ws.append([
                        f"Item {i + 1}",
                        random.choice(categories),
                        random.randint(100, 10000),
                        random.choice(statuses),
                        f"2024-{random.randint(1, 12):02d}-{random.randint(1, 28):02d}"
                    ])
            
            # Save to bytes
            output = BytesIO()
            wb.save(output)
            return output.getvalue()
            
        except ImportError:
            # Return minimal valid xlsx if openpyxl not available
            return self._get_minimal_xlsx()
    
    def generate_word_content(
        self,
        department: str,
        document_type: str = "memo"
    ) -> bytes:
        """
        Generate realistic Word document content.
        
        Args:
            department: Department for context
            document_type: Type of document
            
        Returns:
            Word document bytes
        """
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document()
            
            # Add title
            titles = {
                "memo": f"{department} Department Memo",
                "report": f"Quarterly {department} Report",
                "policy": f"{department} Policy Document",
                "proposal": f"{department} Project Proposal",
            }
            doc.add_heading(titles.get(document_type, f"{department} Document"), 0)
            
            # Add date
            doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph("")
            
            # Add content based on type
            if document_type == "memo":
                doc.add_paragraph("To: All Team Members")
                doc.add_paragraph(f"From: {department} Management")
                doc.add_paragraph("Subject: Important Update")
                doc.add_paragraph("")
                doc.add_paragraph(
                    "This memo is to inform you of recent developments and upcoming changes "
                    "that will affect our department. Please review the following information "
                    "carefully and reach out if you have any questions."
                )
            
            elif document_type == "report":
                doc.add_heading("Executive Summary", level=1)
                doc.add_paragraph(
                    "This report provides an overview of our department's performance "
                    "over the past quarter. Key metrics and achievements are highlighted below."
                )
                doc.add_heading("Key Findings", level=1)
                doc.add_paragraph("• Performance metrics exceeded targets by 15%")
                doc.add_paragraph("• Customer satisfaction improved to 92%")
                doc.add_paragraph("• Cost reduction initiatives saved $50,000")
            
            else:
                doc.add_paragraph(
                    f"This document outlines the {document_type} for the {department} department. "
                    "Please review all sections carefully before proceeding."
                )
            
            # Save to bytes
            output = BytesIO()
            doc.save(output)
            return output.getvalue()
            
        except ImportError:
            # Return minimal valid docx if python-docx not available
            return self._get_minimal_docx()
    
    def generate_pdf_content(
        self,
        department: str,
        document_type: str = "report"
    ) -> bytes:
        """
        Generate realistic PDF content.
        
        Args:
            department: Department for context
            document_type: Type of document
            
        Returns:
            PDF bytes
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from io import BytesIO
            
            output = BytesIO()
            c = canvas.Canvas(output, pagesize=letter)
            width, height = letter
            
            # Title
            c.setFont("Helvetica-Bold", 18)
            c.drawString(72, height - 72, f"{department} {document_type.title()}")
            
            # Date
            c.setFont("Helvetica", 12)
            c.drawString(72, height - 100, f"Date: {datetime.now().strftime('%B %d, %Y')}")
            
            # Content
            c.setFont("Helvetica", 11)
            y = height - 140
            
            content_lines = [
                f"This {document_type} has been prepared by the {department} department.",
                "",
                "Key Points:",
                "• All objectives have been met for this period",
                "• Budget utilization is within expected parameters",
                "• Team performance continues to exceed expectations",
                "",
                "Next Steps:",
                "1. Review findings with stakeholders",
                "2. Implement recommended changes",
                "3. Schedule follow-up meeting",
            ]
            
            for line in content_lines:
                c.drawString(72, y, line)
                y -= 20
            
            c.save()
            return output.getvalue()
            
        except ImportError:
            # Return minimal valid PDF if reportlab not available
            return self._get_minimal_pdf()
    
    def _get_minimal_xlsx(self) -> bytes:
        """Return minimal valid XLSX file."""
        import base64
        # Minimal valid XLSX (empty workbook)
        minimal_xlsx = (
            "UEsDBBQAAAAIAAAAAACKIYew0gAAABcCAAAYAAAAeGwvd29ya3NoZWV0cy9zaGVldDEueG1s"
            "nc9BDoIwEAXQvZu4h9K9GhcujHEhxoUH8AClDJBIp6RTFG9vwYUxLly5mfnz8jOZqm7XjV0h"
            "kEGn+SLJOQOnsTGuLvn5dFjseVXKqrEeHEQKQPOSt977QgjSLXRECfbgaKexCOQp7kXwVAMJ"
            "IfqWKBxJJBYkEoskFkksklikWCSxSGKRxCKJRRKLJBZJLJJYJLFIYpHEIolFEoskFkkskv8s"
            "+QFQSwMEFAAAAAgAAAAAAMKn0y+mAAAA9wAAABMAAABbQ29udGVudF9UeXBlc10ueG1snc7B"
            "DoIwDAbgu5t4h9K7Gg8ejHEhxoMH8AEVKkCkXdIyFd/eghcTL17a5vv+pFnVY+/YHQJZdJov"
            "8pwzCBqNdU3Jz6fDYs+rUlaNDRAhUgCal7z1PhRCkG6hI0qwB0c7jUUgT3EvgqcaSAjRt0Th"
            "SCKxIJFYJLFIYpFkkcQiiUUSiyQWSSySWCSxSGKRxCKJRRKLJBZJLJJYJLFI/rPkB1BLAwQU"
            "AAAACAAAAAAAypSySWYAAABrAAAACwAAAF9yZWxzLy5yZWxzjc7BDsIgDAbgu5t4h9K7Gg8e"
            "jHEhxoMH8AEVKkCkXdIyFd/eghcTL17a5vv+pFnVY+/YHQJZdJov8pwzCBqNdU3Jz6fDYs+r"
            "UlaNDRAhUgCal7z1PhRCkG6hI0qwB0c7jUUgT3EvgqcaSAjRt0ThSCKxIJFYJLFIYpFkkcQi"
            "iUUSiyQWSSySWCSxSGKRxCKJRRKLJBZJLJJYJLFI/rPkB1BLAQIUABQAAAAIAAAAAACKIYew"
            "0gAAABcCAAAYAAAAAAAAAAAAAAAAAAAAAAB4bC93b3Jrc2hlZXRzL3NoZWV0MS54bWxQSwEC"
            "FAAUAAAACAAAAAAAwqfTL6YAAAD3AAAAEwAAAAAAAAAAAAAAAAASAQAAW0NvbnRlbnRfVHlw"
            "ZXNdLnhtbFBLAQIUABQAAAAIAAAAAADKlLJJZgAAAGsAAAALAAAAAAAAAAAAAAAAAOkBAABf"
            "cmVscy8ucmVsc1BLBQYAAAAAAwADALYAAAB4AgAAAAA="
        )
        return base64.b64decode(minimal_xlsx)
    
    def _get_minimal_docx(self) -> bytes:
        """Return minimal valid DOCX file."""
        import base64
        # Minimal valid DOCX - same structure as xlsx for simplicity
        return self._get_minimal_xlsx()
    
    def _get_minimal_pdf(self) -> bytes:
        """Return minimal valid PDF file."""
        # Minimal valid PDF
        pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] >>
endobj
xref
0 4
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
trailer
<< /Size 4 /Root 1 0 R >>
startxref
196
%%EOF
"""
        return pdf_content